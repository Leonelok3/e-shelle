"""
Views pour le module Lebenslauf (CV allemand genere par IA).
"""
import logging
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.conf import settings

from .models import GermanCVProfile, CVExperience, CVEducation, CVLanguage, GeneratedLebenslauf
from .forms import CVProfileForm, CVExperienceForm, CVEducationForm, CVLanguageForm

log = logging.getLogger(__name__)


# ── Prompt systeme IA ─────────────────────────────────────────────────────────
LEBENSLAUF_SYSTEM_PROMPT = """
Tu es un expert RH specialise dans les recrutements Ausbildung en Allemagne.
Tu aides des candidats africains a rediger un Lebenslauf (CV allemand) parfait.

REGLES ABSOLUES DU LEBENSLAUF ALLEMAND :
1. Photo professionnelle (zone reservee si pas de photo)
2. Coordonnees completes : nom, adresse, telephone, email
3. Date de naissance au format DD.MM.YYYY
4. Nationalite mentionnee clairement
5. Parcours professionnel en ordre CHRONOLOGIQUE INVERSE (plus recent en premier)
6. Dates au format MM.YYYY
7. Pas de rubrique "Objectif professionnel"
8. Rubriques standards : Berufserfahrung / Ausbildung / Kenntnisse / Sprachen / Interessen
9. Maximum 2 pages, format DIN A4
10. Langue : ALLEMAND uniquement (sauf la nationalite)

ADAPTE le contenu a l'offre d'Ausbildung ciblee.
Mets en valeur les competences pertinentes pour ce poste specifique.
Si l'experience est dans un secteur different, montre le transfert de competences.

Reponds UNIQUEMENT avec du HTML propre et bien structure (pas de markdown),
utilisant des classes CSS : .lebenslauf, .section-title, .entry, .entry-date,
.entry-content, .skills-grid, .skill-item
"""


@login_required
def dashboard(request):
    """Hub personnel : profil, CV generes, liens rapides."""
    try:
        profile = GermanCVProfile.objects.get(user=request.user)
    except GermanCVProfile.DoesNotExist:
        profile = None

    experiences  = CVExperience.objects.filter(user=request.user)
    educations   = CVEducation.objects.filter(user=request.user)
    languages    = CVLanguage.objects.filter(user=request.user)
    generated    = GeneratedLebenslauf.objects.filter(user=request.user)[:5]

    profile_complete = (
        profile is not None and
        bool(profile.first_name) and
        experiences.exists()
    )

    context = {
        "profile":          profile,
        "experiences":      experiences,
        "educations":       educations,
        "languages":        languages,
        "generated":        generated,
        "profile_complete": profile_complete,
    }
    return render(request, "lebenslauf/dashboard.html", context)


@login_required
def edit_profile(request):
    """Modifier les infos personnelles du candidat."""
    try:
        instance = GermanCVProfile.objects.get(user=request.user)
    except GermanCVProfile.DoesNotExist:
        instance = None

    if request.method == "POST":
        form = CVProfileForm(request.POST, request.FILES, instance=instance)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Profil mis a jour avec succes !")
            return redirect("lebenslauf:dashboard")
    else:
        # Pre-remplir avec les infos Django User si nouveau profil
        initial = {}
        if not instance:
            u = request.user
            initial = {
                "first_name": u.first_name,
                "last_name":  u.last_name,
                "email":      u.email,
            }
        form = CVProfileForm(instance=instance, initial=initial)

    return render(request, "lebenslauf/edit_profile.html", {"form": form})


@login_required
def manage_experiences(request):
    """Ajouter / modifier les experiences professionnelles."""
    if request.method == "POST":
        form = CVExperienceForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Experience ajoutee !")
            return redirect("lebenslauf:dashboard")
    else:
        form = CVExperienceForm()

    experiences = CVExperience.objects.filter(user=request.user)
    return render(request, "lebenslauf/manage_experiences.html", {
        "form": form, "experiences": experiences
    })


@login_required
def delete_experience(request, pk):
    exp = get_object_or_404(CVExperience, pk=pk, user=request.user)
    exp.delete()
    return redirect("lebenslauf:dashboard")


@login_required
def manage_education(request):
    if request.method == "POST":
        form = CVEducationForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Formation ajoutee !")
            return redirect("lebenslauf:dashboard")
    else:
        form = CVEducationForm()

    educations = CVEducation.objects.filter(user=request.user)
    return render(request, "lebenslauf/manage_education.html", {
        "form": form, "educations": educations
    })


@login_required
def manage_languages(request):
    if request.method == "POST":
        form = CVLanguageForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Langue ajoutee !")
            return redirect("lebenslauf:dashboard")
    else:
        form = CVLanguageForm()

    languages = CVLanguage.objects.filter(user=request.user)
    return render(request, "lebenslauf/manage_languages.html", {
        "form": form, "languages": languages
    })


@login_required
def generate_lebenslauf(request, offer_pk=None):
    """
    Genere un Lebenslauf adapte a une offre specifique (ou candidature spontanee).
    Appelle GPT-4o et sauvegarde le HTML.
    """
    from germany_opportunities.models import AusbildungOffer

    offer = None
    if offer_pk:
        offer = get_object_or_404(AusbildungOffer, pk=offer_pk)

    try:
        profile = GermanCVProfile.objects.get(user=request.user)
    except GermanCVProfile.DoesNotExist:
        messages.error(request, "Completez d'abord votre profil avant de generer un Lebenslauf.")
        return redirect("lebenslauf:edit_profile")

    if request.method == "POST":
        custom_offer_title   = request.POST.get("custom_offer_title", "")
        custom_offer_company = request.POST.get("custom_offer_company", "")

        experiences = CVExperience.objects.filter(user=request.user)
        educations  = CVEducation.objects.filter(user=request.user)
        languages   = CVLanguage.objects.filter(user=request.user)

        # Construire le contexte candidat
        candidate_context = _build_candidate_context(
            profile, experiences, educations, languages
        )

        # Construire le contexte offre
        if offer:
            offer_context = (
                f"Poste : {offer.title}\n"
                f"Entreprise : {offer.company}\n"
                f"Ville : {offer.city}\n"
                f"Secteur : {offer.get_sector_display()}\n"
                f"Description : {offer.description[:600]}"
            )
        elif custom_offer_title:
            offer_context = (
                f"Poste : {custom_offer_title}\n"
                f"Entreprise : {custom_offer_company}"
            )
        else:
            offer_context = "Candidature spontanee dans le secteur : " + profile.target_sector

        html_content, cover_letter = _call_ai_generate(candidate_context, offer_context)

        if html_content:
            generated = GeneratedLebenslauf.objects.create(
                user=request.user,
                offer=offer,
                custom_offer_title=custom_offer_title,
                custom_offer_company=custom_offer_company,
                content_html=html_content,
                ai_cover_letter=cover_letter,
            )
            return redirect("lebenslauf:view_lebenslauf", pk=generated.pk)
        else:
            messages.error(request, "Erreur lors de la generation. Reessayez.")

    context = {
        "profile": profile,
        "offer":   offer,
    }
    return render(request, "lebenslauf/generate.html", context)


@login_required
def view_lebenslauf(request, pk):
    """Apercu du Lebenslauf genere."""
    lv = get_object_or_404(GeneratedLebenslauf, pk=pk, user=request.user)
    return render(request, "lebenslauf/view_lebenslauf.html", {"lebenslauf": lv})


@login_required
def download_lebenslauf(request, pk):
    """Telechargement du Lebenslauf en HTML (imprimable)."""
    lv = get_object_or_404(GeneratedLebenslauf, pk=pk, user=request.user)
    response = HttpResponse(lv.content_html, content_type="text/html; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="lebenslauf_{pk}.html"'
    return response


class DocxHTMLParser(HTMLParser):
    def __init__(self, cell_or_doc, default_text_color, heading_color, accent_color=None):
        super().__init__()
        self.container = cell_or_doc
        self.default_text_color = default_text_color
        self.heading_color = heading_color
        self.accent_color = accent_color or heading_color
        self.current_tag = None
        self.current_paragraph = None
        
    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ('h1', 'h2', 'h3', 'h4'):
            self.current_paragraph = self.container.add_paragraph()
            self.current_paragraph.paragraph_format.space_before = Pt(8)
            self.current_paragraph.paragraph_format.space_after = Pt(4)
        elif tag in ('p', 'div'):
            self.current_paragraph = self.container.add_paragraph()
            self.current_paragraph.paragraph_format.space_after = Pt(4)
        elif tag == 'li':
            self.current_paragraph = self.container.add_paragraph(style='List Bullet')
            self.current_paragraph.paragraph_format.space_after = Pt(2)
        elif tag == 'br':
            if self.current_paragraph:
                self.current_paragraph.add_run('\n')
                
    def handle_endtag(self, tag):
        self.current_tag = None
        
    def handle_data(self, data):
        text = data.strip()
        if not text:
            return
        if self.current_tag in ('h1', 'h2', 'h3', 'h4'):
            run = self.current_paragraph.add_run(text)
            run.bold = True
            if self.current_tag == 'h1':
                run.font.size = Pt(15)
                run.font.color.rgb = self.heading_color
            elif self.current_tag == 'h2':
                run.font.size = Pt(11.5)
                run.font.color.rgb = self.accent_color
            else:
                run.font.size = Pt(10)
                run.font.color.rgb = self.default_text_color
        elif self.current_paragraph:
            run = self.current_paragraph.add_run(" " + text if self.current_paragraph.runs else text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = self.default_text_color
            if self.current_tag == 'strong':
                run.bold = True


@login_required
def download_lebenslauf_docx(request, pk):
    """Téléchargement du CV et de la Lettre de motivation au format Word (.docx)."""
    import io
    import re
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    lv = get_object_or_404(GeneratedLebenslauf, pk=pk, user=request.user)
    
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Helper function to shade cells
    def shade_cell(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    # Helper function to set cell padding (in dxa: 20 dxa = 1 point)
    def set_cell_padding(cell, top=200, bottom=200, left=300, right=300):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # 1. Add Cover Letter (Anschreiben) if it exists
    if lv.ai_cover_letter:
        months = ["Januar", "Februar", "März", "April", "Mai", "Juni", "Juli", "August", "September", "Oktober", "November", "Dezember"]
        
        for line in lv.ai_cover_letter.split('\n'):
            cleaned_line = line.strip()
            if not cleaned_line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                continue
                
            is_date_line = False
            if any(m in cleaned_line for m in months) and re.search(r'\d{4}', cleaned_line) and len(cleaned_line) < 50:
                is_date_line = True
                
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            if is_date_line:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.add_run(cleaned_line)
            else:
                parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
            
        doc.add_page_break()

    # 2. Add CV (Lebenslauf) 2-column Table
    table = doc.add_table(rows=1, cols=2)
    # Ensure column widths
    table.columns[0].width = Inches(2.3)
    table.columns[1].width = Inches(4.7)
    
    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    
    # Design Left Spalte (Sidebar)
    shade_cell(left_cell, "2D3748")  # Dark slate gray
    set_cell_padding(left_cell, top=300, bottom=300, left=250, right=250)
    
    # Design Right Spalte (Main)
    shade_cell(right_cell, "FFFFFF")
    set_cell_padding(right_cell, top=300, bottom=300, left=250, right=250)

    # Clean HTML from style/script tags
    clean_html = re.sub(r'<style[^>]*>.*?</style>', '', lv.content_html, flags=re.DOTALL)
    clean_html = re.sub(r'<script[^>]*>.*?</script>', '', clean_html, flags=re.DOTALL)

    # Parse and split columns
    left_html = ""
    right_html = ""
    if "<!-- Rechte Spalte" in clean_html:
        parts = clean_html.split("<!-- Rechte Spalte")
        left_html = parts[0].strip()
        right_html = parts[1].split("-->", 1)[1].strip()
    else:
        left_html = clean_html
        right_html = ""

    # Parse Left Column
    left_parser = DocxHTMLParser(
        left_cell,
        default_text_color=RGBColor(247, 250, 252),  # Light gray text
        heading_color=RGBColor(104, 211, 145),       # Mint Green #68D391
        accent_color=RGBColor(255, 206, 0)           # Gold/Yellow
    )
    left_parser.feed(left_html)

    # Parse Right Column
    if right_html:
        right_parser = DocxHTMLParser(
            right_cell,
            default_text_color=RGBColor(45, 55, 72),     # Dark slate text
            heading_color=RGBColor(45, 55, 72),
            accent_color=RGBColor(227, 0, 15)            # German Red
        )
        right_parser.feed(right_html)

    # Save to memory stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"bewerbung_{pk}.docx"
    response = HttpResponse(
        file_stream.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ── Helpers ───────────────────────────────────────────────────────────────────

def _build_candidate_context(profile, experiences, educations, languages):
    lines = [
        f"Candidat : {profile.full_name}",
        f"Nationalite : {profile.nationality}",
        f"Date de naissance : {profile.date_of_birth.strftime('%d.%m.%Y') if profile.date_of_birth else 'N/A'}",
        f"Email : {profile.email}",
        f"Telephone : {profile.phone}",
        f"Adresse actuelle : {profile.address}",
        f"Niveau d'allemand : {profile.german_level}",
        f"Certificat Goethe obtenu : {'Oui' if profile.goethe_certified else 'Non'}",
        "",
        "=== EXPERIENCES PROFESSIONNELLES ===",
    ]
    for exp in experiences:
        lines += [
            f"- {exp.title} chez {exp.company} ({exp.city}, {exp.country})",
            f"  Periode : {exp.period_display}",
            f"  Description : {exp.description}",
        ]

    lines += ["", "=== FORMATIONS ==="]
    for edu in educations:
        lines += [
            f"- {edu.degree} — {edu.school} ({edu.city}, {edu.country})",
            f"  {edu.start_year} – {edu.end_year or 'en cours'}",
        ]

    lines += ["", "=== LANGUES ==="]
    for lang in languages:
        cert = f" ({lang.certificate})" if lang.certificate else ""
        lines += [f"- {lang.language} : {lang.proficiency}{cert}"]

    return "\n".join(lines)


def _call_ai_generate(candidate_context: str, offer_context: str) -> tuple[str, str]:
    """
    Appelle Gemini pour générer à la fois le Lebenslauf en HTML et la lettre de motivation (Anschreiben)
    en un seul appel pour des raisons de performance et pour éviter le timeout.
    """
    import datetime
    from ai_engine.services.llm_service import call_llm

    now = datetime.datetime.now()
    months_de = {
        1: "Januar", 2: "Februar", 3: "März", 4: "April", 5: "Mai", 6: "Juni",
        7: "Juli", 8: "August", 9: "September", 10: "Oktober", 11: "November", 12: "Dezember"
    }
    today_german = f"{now.day}. {months_de[now.month]} {now.year}"

    system_prompt = (
        "Du bist ein renommierter deutscher HR-Experte, spezialisiert auf das deutsche Bewerbungsverfahren.\n"
        "Deine Aufgabe ist es, für den Kandidaten zwei Dokumente in fehlerfreiem, professionellem Deutsch zu erstellen:\n"
        "1. Einen modernen, professionellen Lebenslauf (CV) im zweiseitigen/zweispaltigen Tabellen-Layout (HTML).\n"
        f"2. Ein überzeugendes Anschreiben (Bewerbungsschreiben) auf den heutigen Tag ({today_german}) datiert (Text).\n\n"
        "Übersetze alle Informationen des Kandidaten auf Deutsch und formuliere seine Erfahrungen so um, dass die "
        "übertragbaren Kompetenzen (Pünktlichkeit, Organisation, Empathie) für das Ziel-Ausbildungsangebot im Vordergrund stehen.\n\n"
        "REGLEN FÜR DAS HTML-LAYOUT DES LEBENSLAUFS:\n"
        "- Generiere einen Container mit zwei Spalten: linke Spalte dunkelgrau (#2D3748) für persönliche Daten & Sprachen, "
        "rechte Spalte weiß für Lebenslauf, Berufserfahrung, Ausbildung, Kenntnisse.\n"
        "- Verwende Flexbox/Grid mit Inlinestilen.\n"
        "- Nutze das deutsche Datumsformat (MM.YYYY).\n\n"
        "REGLEN FÜR DAS ANSCHREIBEN (UNBEDINGT EINHALTEN):\n"
        f"- Verwende als Ort und Datum genau dieses Format auf einer eigenen Zeile: '[Stadt des Kandidaten], {today_german}'.\n"
        "- Die Betreffzeile (z.B. 'Bewerbung um einen Ausbildungsplatz als...') MUSS fett gedruckt sein unter Verwendung von Markdown ** (z.B. **Bewerbung um einen Ausbildungsplatz als...**).\n"
        "- Das gesamte Anschreiben MUSS auf genau eine einzige DIN-A4-Seite passen (maximal 200 bis 230 Wörter für den Hauptteil). Sei präzise und komm direkt auf den Punkt. Längere Briefe werden in Deutschland abgelehnt.\n"
        "- Beende das Anschreiben immer mit der Grußformel 'Mit freundlichen Grüßen' und darunter dem Namen des Kandidaten als Unterschrift.\n\n"
        "FORMATIERUNG DER ANTWORT:\n"
        "Gib zuerst den reinen HTML-Code des Lebenslaufs aus. Kein Markdown (kein ```html).\n"
        "Füge danach die genaue Trennzeile ein:\n"
        "=== ANSCHREIBEN_START ===\n"
        "Gib danach den reinen Text des Anschreibens aus."
    )

    user_prompt = (
        f"KANDIDATENPROFIL:\n{candidate_context}\n\n"
        f"ZIELSTELLENANGEBOT:\n{offer_context}\n\n"
        "Erstelle jetzt den Lebenslauf (HTML) und das Anschreiben (Text)."
    )

    try:
        response = call_llm(system_prompt, user_prompt)
        if not response:
            return "", ""

        # Clean markdown code blocks if any
        response = response.replace("```html", "").replace("```", "").strip()

        if "=== ANSCHREIBEN_START ===" in response:
            parts = response.split("=== ANSCHREIBEN_START ===")
            html_content = parts[0].strip()
            cover_letter = parts[1].strip()
        else:
            # Fallback if delimiter not found
            html_content = response
            cover_letter = ""

        return html_content, cover_letter
    except Exception as exc:
        log.error(f"Lebenslauf AI generation error: {exc}")
        return "", ""
