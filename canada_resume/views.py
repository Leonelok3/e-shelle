import logging
import io
import re
import datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.contrib import messages
from django.urls import reverse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import CanadaCVProfile, CanadaCVExperience, CanadaCVEducation, CanadaCVLanguage, GeneratedCanadaResume, CanadaImmigrationProfile, CanadaCVProject, CanadaResource
from .forms import CanadaCVProfileForm, CanadaCVExperienceForm, CanadaCVEducationForm, CanadaCVLanguageForm, CanadaImmigrationProfileForm, CanadaCVProjectForm
from jobs.models import CanadaJobOffer

log = logging.getLogger(__name__)

def check_user_has_paid_edu_subscription(user) -> bool:
    if not user.is_authenticated:
        return False
    if user.is_superuser or user.is_staff:
        return True
    if hasattr(user, "profile") and user.profile.plan in ["pro", "enterprise"]:
        return True
    try:
        from accounts.models import AppSubscription
        sub = AppSubscription.get_active_for_user(user, "prep")
        if sub and sub.is_active and not sub.plan.is_free:
            return True
    except Exception:
        pass
    return False


@login_required
def dashboard(request):
    try:
        profile = CanadaCVProfile.objects.get(user=request.user)
    except CanadaCVProfile.DoesNotExist:
        profile = None

    experiences = CanadaCVExperience.objects.filter(user=request.user)
    educations = CanadaCVEducation.objects.filter(user=request.user)
    languages = CanadaCVLanguage.objects.filter(user=request.user)
    projects = CanadaCVProject.objects.filter(user=request.user)
    generated = GeneratedCanadaResume.objects.filter(user=request.user)[:5]

    profile_complete = (
        profile is not None and
        bool(profile.first_name) and
        experiences.exists()
    )

    context = {
        "profile": profile,
        "experiences": experiences,
        "educations": educations,
        "languages": languages,
        "projects": projects,
        "generated": generated,
        "profile_complete": profile_complete,
        "is_pro": check_user_has_paid_edu_subscription(request.user),
    }
    return render(request, "canada_resume/dashboard.html", context)


@login_required
def edit_profile(request):
    try:
        instance = CanadaCVProfile.objects.get(user=request.user)
    except CanadaCVProfile.DoesNotExist:
        instance = None

    if request.method == "POST":
        form = CanadaCVProfileForm(request.POST, instance=instance)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Votre profil a été enregistré avec succès !")
            return redirect("canada_resume:dashboard")
    else:
        initial = {}
        if not instance:
            u = request.user
            initial = {
                "first_name": u.first_name,
                "last_name": u.last_name,
                "email": u.email,
            }
        form = CanadaCVProfileForm(instance=instance, initial=initial)

    return render(request, "canada_resume/edit_profile.html", {"form": form})


@login_required
def manage_experiences(request):
    if request.method == "POST":
        form = CanadaCVExperienceForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Expérience professionnelle ajoutée !")
            return redirect("canada_resume:dashboard")
    else:
        form = CanadaCVExperienceForm()

    experiences = CanadaCVExperience.objects.filter(user=request.user)
    return render(request, "canada_resume/manage_experiences.html", {
        "form": form, "experiences": experiences
    })


@login_required
def delete_experience(request, pk):
    exp = get_object_or_404(CanadaCVExperience, pk=pk, user=request.user)
    exp.delete()
    messages.success(request, "Expérience supprimée.")
    return redirect("canada_resume:dashboard")


@login_required
def manage_education(request):
    if request.method == "POST":
        form = CanadaCVEducationForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Formation académique ajoutée !")
            return redirect("canada_resume:dashboard")
    else:
        form = CanadaCVEducationForm()

    educations = CanadaCVEducation.objects.filter(user=request.user)
    return render(request, "canada_resume/manage_education.html", {
        "form": form, "educations": educations
    })


@login_required
def delete_education(request, pk):
    edu = get_object_or_404(CanadaCVEducation, pk=pk, user=request.user)
    edu.delete()
    messages.success(request, "Formation supprimée.")
    return redirect("canada_resume:dashboard")


@login_required
def manage_languages(request):
    if request.method == "POST":
        form = CanadaCVLanguageForm(request.POST)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Compétence linguistique ajoutée !")
            return redirect("canada_resume:dashboard")
    else:
        form = CanadaCVLanguageForm()

    languages = CanadaCVLanguage.objects.filter(user=request.user)
    return render(request, "canada_resume/manage_languages.html", {
        "form": form, "languages": languages
    })


@login_required
def delete_language(request, pk):
    lang = get_object_or_404(CanadaCVLanguage, pk=pk, user=request.user)
    lang.delete()
    messages.success(request, "Compétence linguistique supprimée.")
    return redirect("canada_resume:dashboard")


@login_required
def generate_resume(request, offer_pk=None):
    offer = None
    if offer_pk:
        offer = get_object_or_404(CanadaJobOffer, pk=offer_pk)

    try:
        profile = CanadaCVProfile.objects.get(user=request.user)
    except CanadaCVProfile.DoesNotExist:
        messages.error(request, "Veuillez d'abord compléter votre profil avant de pouvoir générer un CV.")
        return redirect("canada_resume:edit_profile")

    if request.method == "POST":
        custom_offer_title = request.POST.get("custom_offer_title", "")
        custom_offer_company = request.POST.get("custom_offer_company", "")
        lang_choice = request.POST.get("language", "fr")

        experiences = CanadaCVExperience.objects.filter(user=request.user)
        educations = CanadaCVEducation.objects.filter(user=request.user)
        languages = CanadaCVLanguage.objects.filter(user=request.user)

        candidate_context = _build_candidate_context(profile, experiences, educations, languages)

        if offer:
            offer_context = (
                f"Poste : {offer.title}\n"
                f"Entreprise : {offer.company}\n"
                f"Lieu : {offer.city}, {offer.province}\n"
                f"Description : {offer.description}"
            )
        elif custom_offer_title:
            offer_context = (
                f"Poste : {custom_offer_title}\n"
                f"Entreprise : {custom_offer_company}"
            )
        else:
            offer_context = f"Candidature spontanée dans le secteur : {profile.target_sector}"

        html_content, cover_letter = _call_ai_generate_canada(candidate_context, offer_context, lang_choice)

        if html_content:
            generated = GeneratedCanadaResume.objects.create(
                user=request.user,
                offer=offer,
                custom_offer_title=custom_offer_title,
                custom_offer_company=custom_offer_company,
                language=lang_choice,
                content_html=html_content,
                ai_cover_letter=cover_letter,
            )
            return redirect("canada_resume:view_resume", pk=generated.pk)
        else:
            messages.error(request, "Une erreur est survenue lors de la génération. Veuillez réessayer.")

    context = {
        "profile": profile,
        "offer": offer,
    }
    return render(request, "canada_resume/generate.html", context)


@login_required
def view_resume(request, pk):
    resume = get_object_or_404(GeneratedCanadaResume, pk=pk, user=request.user)
    return render(request, "canada_resume/view_resume.html", {
        "resume": resume,
        "is_pro": check_user_has_paid_edu_subscription(request.user),
    })


@login_required
def download_resume_docx(request, pk):
    if not check_user_has_paid_edu_subscription(request.user):
        messages.warning(request, "Le téléchargement du CV au format Word (.docx) est réservé aux abonnés Premium.")
        return redirect("canada_resume:dashboard")

    resume = get_object_or_404(GeneratedCanadaResume, pk=pk, user=request.user)
    doc = Document()

    # Configuration des marges (Format canadien compact 0.5 in)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Nettoyage de l'HTML (Sanitisation des sauts de ligne excessifs et paragraphes vides)
    clean_html = re.sub(r'<style[^>]*>.*?</style>', '', resume.content_html, flags=re.DOTALL)
    clean_html = re.sub(r'<script[^>]*>.*?</script>', '', clean_html, flags=re.DOTALL)
    clean_html = re.sub(r'(?:<br\s*/?>\s*)+', '<br/>', clean_html)
    
    # Strip layout divs entirely so we only rely on p/ul/li/h tags for paragraphs, preventing extra spacing
    clean_html = re.sub(r'</?div[^>]*>', '', clean_html)
    
    clean_html = re.sub(r'<p\b[^>]*>\s*(?:&nbsp;|<br\s*/?>|\s)*</p>', '', clean_html)
    
    # Remplacement simple des balises HTML pour docx
    from html.parser import HTMLParser

    class SimpleDocxParser(HTMLParser):
        def __init__(self, document):
            super().__init__()
            self.doc = document
            self.current_paragraph = None
            self.current_tag = None

        def handle_starttag(self, tag, attrs):
            self.current_tag = tag
            if tag in ('h1', 'h2', 'h3'):
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.paragraph_format.space_before = Pt(8)
                self.current_paragraph.paragraph_format.space_after = Pt(2)
                self.current_paragraph.paragraph_format.keep_with_next = True
            elif tag in ('p', 'div'):
                self.current_paragraph = self.doc.add_paragraph()
                self.current_paragraph.paragraph_format.space_before = Pt(0)
                self.current_paragraph.paragraph_format.space_after = Pt(2)
            elif tag == 'li':
                self.current_paragraph = self.doc.add_paragraph(style='List Bullet')
                self.current_paragraph.paragraph_format.space_before = Pt(0)
                self.current_paragraph.paragraph_format.space_after = Pt(1)
                self.current_paragraph.paragraph_format.left_indent = Inches(0.25)
            elif tag == 'br':
                if self.current_paragraph:
                    self.current_paragraph.add_run('\n')

        def handle_endtag(self, tag):
            self.current_tag = None

        def handle_data(self, data):
            text = data.strip()
            if not text:
                return
            if self.current_tag in ('h1', 'h2', 'h3'):
                run = self.current_paragraph.add_run(text)
                run.bold = True
                if self.current_tag == 'h1':
                    run.font.size = Pt(16)
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(16, 43, 78) # Dark Blue #102B4E
            elif self.current_paragraph:
                run = self.current_paragraph.add_run(" " + text if self.current_paragraph.runs else text)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
                if self.current_tag == 'strong':
                    run.bold = True

    parser = SimpleDocxParser(doc)
    parser.feed(clean_html)

    # Génération du flux
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"cv_canada_{pk}.docx"
    response = HttpResponse(
        file_stream.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def download_cover_letter_docx(request, pk):
    if not check_user_has_paid_edu_subscription(request.user):
        messages.warning(request, "Le téléchargement de la lettre de motivation est réservé aux abonnés Premium.")
        return redirect("canada_resume:dashboard")

    resume = get_object_or_404(GeneratedCanadaResume, pk=pk, user=request.user)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    if resume.ai_cover_letter:
        for line in resume.ai_cover_letter.split('\n'):
            cleaned_line = line.strip()
            if not cleaned_line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            
            parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"lettre_motivation_canada_{pk}.docx"
    response = HttpResponse(
        file_stream.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _build_candidate_context(profile, experiences, educations, languages):
    lines = [
        f"Nom complet : {profile.full_name}",
        f"Email : {profile.email}",
        f"Téléphone : {profile.phone}",
        f"Adresse : {profile.address}",
        f"LinkedIn : {profile.linkedin}",
        f"Secteur d'activité ciblé : {profile.target_sector}",
        f"Provinces acceptées au Canada : {profile.target_provinces}",
        "",
        "=== EXPÉRIENCES PROFESSIONNELLES ===",
    ]
    for exp in experiences:
        lines += [
            f"- {exp.title} chez {exp.company} ({exp.city}, {exp.province_country})",
            f"  Période : {exp.period_display_fr}",
            f"  Description & Réalisations : {exp.description}",
        ]

    lines += ["", "=== FORMATIONS ==="]
    for edu in educations:
        lines += [
            f"- {edu.degree} — {edu.school} ({edu.city}, {edu.province_country})",
            f"  Années : {edu.start_year} – {edu.end_year or 'Présent'}",
        ]

    lines += ["", "=== LANGUES ==="]
    for lang in languages:
        cert = f" ({lang.certificate})" if lang.certificate else ""
        lines += [f"- {lang.language} : {lang.proficiency}{cert}"]

    return "\n".join(lines)


def _call_ai_generate_canada(candidate_context: str, offer_context: str, lang_choice: str) -> tuple[str, str]:
    import datetime
    from ai_engine.services.llm_service import call_llm

    today_str = datetime.date.today().strftime("%d/%m/%Y")

    if lang_choice == "en":
        system_prompt = (
            "You are an expert HR recruiter specializing in Canadian immigration and ATS-friendly resumes.\n"
            "Your task is to generate two professional documents in English for the candidate:\n"
            "1. An ATS-friendly Canadian resume (HTML format).\n"
            "2. A compelling Canadian cover letter (Text format).\n\n"
            "CRITICAL CANADIAN RESUME RULES:\n"
            "- NO photo. Never include or leave a placeholder for a photo.\n"
            "- NO birth date, age, marital status, gender, or nationality. Including these is illegal in Canada and causes ATS rejection.\n"
            "- Use clean single-column format for the HTML layout with standard headings: SUMMARY, EXPERIENCE, EDUCATION, SKILLS.\n"
            "- ATS-FRIENDLY & COMPACT STYLE RULES:\n"
            "  * Wrap the entire resume in: <div style=\"font-family:'Calibri', 'Arial', sans-serif; line-height:1.25; font-size:10.5pt; color:#1e293b; max-width:800px; margin:0 auto; padding:0;\">\n"
            "  * Name Header: Use <h1 style=\"text-align:center; font-size:18pt; font-weight:bold; color:#102B4E; margin:0 0 4px 0; padding:0;\">[Full Name]</h1>\n"
            "  * Contact Info: Use <p style=\"text-align:center; font-size:9.5pt; color:#475569; margin:0 0 12px 0; padding:0;\">[Email] | [Phone] | [Address] | [LinkedIn]</p>\n"
            "  * Section Headings (SUMMARY, EXPERIENCE, etc.): Use <h2 style=\"font-size:12pt; font-weight:bold; color:#102B4E; border-bottom:1.5px solid #102B4E; margin:14px 0 6px 0; padding:0 0 2px 0; text-transform:uppercase;\">[Heading]</h2>\n"
            "  * Job Title / Degree entries: Use <div style=\"display:flex; justify-content:space-between; font-weight:bold; font-size:10pt; color:#1e293b; margin:6px 0 2px 0;\"><span>[Title/Degree] at [Company/School]</span><span>[Period/Years]</span></div>\n"
            "  * Location: Use <p style=\"font-style:italic; font-size:9.5pt; color:#475569; margin:0 0 4px 0;\">[City, Province/Country]</p>\n"
            "  * Achievements list: Use <ul style=\"padding-left:16px; margin:2px 0 6px 0;\"> and <li style=\"margin-bottom:2px; font-size:9.5pt; color:#334155;\">[Bullet]</li>\n"
            "  * Skills section: Group skills by category, and render each skill as a compact inline-block badge block. For example:\n"
            "    <div style=\"margin-bottom:6px;\">\n"
            "      <strong style=\"font-size:9.5pt; color:#102B4E;\">[Category Name]:</strong>\n"
            "      <div style=\"margin-top:4px;\">\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">[Skill 1]</span>\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">[Skill 2]</span>\n"
            "      </div>\n"
            "    </div>\n"
            "  * Prohibited: Do NOT use consecutive <br> tags, empty paragraphs like <p>&nbsp;</p> or large blank margin/padding. Keep everything compact so it fits beautifully on 1 page.\n\n"
            "CRITICAL COVER LETTER RULES:\n"
            "- The Cover Letter must start with candidate coordinates, date, employer details, subject line in bold (using **), body text, and closing (Sincerely, [Candidate Name]).\n\n"
            "RESPONSE FORMAT:\n"
            "Output the raw HTML for the resume first. Do not wrap in ```html block.\n"
            "Then output the divider exactly:\n"
            "=== ANSCHREIBEN_START ===\n"
            "Then output the plain text cover letter."
        )
    elif lang_choice == "bi":
        system_prompt = (
            "You are an expert HR recruiter specializing in Canadian immigration and bilingual (English/French) ATS-friendly resumes.\n"
            "Your task is to generate two professional documents in a side-by-side or slash-divided bilingual (FR/EN) format for the candidate:\n"
            "1. A bilingual ATS-friendly Canadian resume (HTML format).\n"
            "2. A bilingual Canadian cover letter (Text format).\n\n"
            "CRITICAL BILINGUAL RESUME RULES:\n"
            "- NO photo. Never include or leave a placeholder for a photo.\n"
            "- NO birth date, age, marital status, gender, or nationality.\n"
            "- Use clean single-column format for the HTML layout with bilingual headings separated by a slash: PROFIL / SUMMARY, EXPÉRIENCE PROFESSIONNELLE / EXPERIENCE, FORMATION / EDUCATION, COMPÉTENCES / SKILLS, LANGUES / LANGUAGES.\n"
            "- BILINGUAL WRITING STYLES:\n"
            "  * Wrap the entire resume in: <div style=\"font-family:'Calibri', 'Arial', sans-serif; line-height:1.25; font-size:10.5pt; color:#1e293b; max-width:800px; margin:0 auto; padding:0;\">\n"
            "  * Name Header: Use <h1 style=\"text-align:center; font-size:18pt; font-weight:bold; color:#102B4E; margin:0 0 4px 0; padding:0;\">[Full Name]</h1>\n"
            "  * Contact Info: Use <p style=\"text-align:center; font-size:9.5pt; color:#475569; margin:0 0 12px 0; padding:0;\">[Email] | [Phone] | [Address] | [LinkedIn]</p>\n"
            "  * Section Headings: e.g. <h2 style=\"font-size:12pt; font-weight:bold; color:#102B4E; border-bottom:1.5px solid #102B4E; margin:14px 0 6px 0; padding:0 0 2px 0; text-transform:uppercase;\">PROFIL / SUMMARY</h2>\n"
            "  * Job Title / Degree entries: e.g. <div style=\"display:flex; justify-content:space-between; font-weight:bold; font-size:10pt; color:#1e293b; margin:6px 0 2px 0;\"><span>[Poste en FR] / [Job Title in EN] chez/at [Company]</span><span>[Period/Years]</span></div>\n"
            "  * Location/City: e.g. <p style=\"font-style:italic; font-size:9.5pt; color:#475569; margin:0 0 4px 0;\">[Douala, Cameroun / Douala, Cameroon]</p>\n"
            "  * Descriptions & Achievements: For each achievement bullet or description, output the French version first, followed immediately by the English version in italics, e.g.:\n"
            "    <li style=\"margin-bottom:4px; font-size:9.5pt; color:#334155;\">Géré le contenu et la fonctionnalité du site web.<br><span style=\"font-style:italic; color:#64748b;\">Managed website content and functionality.</span></li>\n"
            "  * Skills section: Group skills, and render each skill as a compact inline-block badge block. For example:\n"
            "    <div style=\"margin-bottom:6px;\">\n"
            "      <strong style=\"font-size:9.5pt; color:#102B4E;\">Développement Web / Web Development:</strong>\n"
            "      <div style=\"margin-top:4px;\">\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">Django</span>\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">React</span>\n"
            "      </div>\n"
            "    </div>\n"
            "  * Prohibited: Do NOT use consecutive <br> tags, empty paragraphs like <p>&nbsp;</p> or large blank margin/padding. Keep everything compact so it fits beautifully on 1 page.\n\n"
            "CRITICAL BILINGUAL COVER LETTER RULES:\n"
            "- Start with candidate coordinates, date, employer details.\n"
            "- Ligne d'Objet / Subject Line: Object / Subject: **[Objet en FR] / [Subject in EN]**\n"
            "- Render the full letter body first in French, and then render the full letter body in English in italics right below it.\n"
            "- End with both closing salutations (e.g. Cordialement, / Sincerely, [Candidate Name]).\n\n"
            "RESPONSE FORMAT:\n"
            "Output the raw HTML for the resume first. Do not wrap in ```html block.\n"
            "Then output the divider exactly:\n"
            "=== ANSCHREIBEN_START ===\n"
            "Then output the plain text cover letter."
        )
    else:
        system_prompt = (
            "Tu es un expert en recrutement canadien spécialisé dans l'immigration et la rédaction de CV compatibles ATS.\n"
            "Ta tâche est de générer deux documents professionnels en Français canadien pour le candidat :\n"
            "1. Un CV au format canadien compatible ATS (format HTML).\n"
            "2. Une lettre de présentation / motivation convaincante (format Texte).\n\n"
            "RÈGLES CRITIQUES DU CV CANADIEN :\n"
            "- AUCUNE photo. Ne jamais inclure de photo ni de zone réservée pour une photo.\n"
            "- AUCUNE mention de l'âge, date de naissance, genre, statut marital ou nationalité (Interdit par la loi).\n"
            "- CV rédigé en colonne simple avec des rubriques standards : PROFIL, EXPÉRIENCE PROFESSIONNELLE, FORMATION, COMPÉTENCES.\n"
            "- RÈGLES DE STYLE ATS COMPACT ET PROFESSIONNEL :\n"
            "  * Enveloppe tout le CV dans : <div style=\"font-family:'Calibri', 'Arial', sans-serif; line-height:1.25; font-size:10.5pt; color:#1e293b; max-width:800px; margin:0 auto; padding:0;\">\n"
            "  * En-tête Nom : Utilise <h1 style=\"text-align:center; font-size:18pt; font-weight:bold; color:#102B4E; margin:0 0 4px 0; padding:0;\">[Nom Complet]</h1>\n"
            "  * Coordonnées : Utilise <p style=\"text-align:center; font-size:9.5pt; color:#475569; margin:0 0 12px 0; padding:0;\">[Email] | [Téléphone] | [Adresse] | [LinkedIn]</p>\n"
            "  * Titres de sections (PROFIL, EXPÉRIENCE PROFESSIONNELLE, etc.) : Utilise <h2 style=\"font-size:12pt; font-weight:bold; color:#102B4E; border-bottom:1.5px solid #102B4E; margin:14px 0 6px 0; padding:0 0 2px 0; text-transform:uppercase;\">[Titre]</h2>\n"
            "  * Postes / Diplômes : Utilise <div style=\"display:flex; justify-content:space-between; font-weight:bold; font-size:10pt; color:#1e293b; margin:6px 0 2px 0;\"><span>[Poste/Diplôme] chez [Entreprise/École]</span><span>[Période/Années]</span></div>\n"
            "  * Ville et Pays : Utilise <p style=\"font-style:italic; font-size:9.5pt; color:#475569; margin:0 0 4px 0;\">[Ville, Province/Pays]</p>\n"
            "  * Listes à puces : Utilise <ul style=\"padding-left:16px; margin:2px 0 6px 0;\"> et <li style=\"margin-bottom:2px; font-size:9.5pt; color:#334155;\">[Puce]</li>\n"
            "  * Section Compétences : Formate chaque catégorie sous forme de blocs/badges (spans inline-block) comme ceci :\n"
            "    <div style=\"margin-bottom:6px;\">\n"
            "      <strong style=\"font-size:9.5pt; color:#102B4E;\">[Nom de la Catégorie] :</strong>\n"
            "      <div style=\"margin-top:4px;\">\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">[Compétence 1]</span>\n"
            "        <span style=\"display:inline-block; background:#f1f5f9; color:#1e293b; padding:2px 6px; border-radius:4px; font-size:9pt; border:1px solid #cbd5e1; margin-right:4px; margin-bottom:4px; font-weight:500;\">[Compétence 2]</span>\n"
            "      </div>\n"
            "    </div>\n"
            "  * Interdictions : Ne pas utiliser de balises <br> consécutives, de paragraphes vides comme <p>&nbsp;</p> ou de marges géantes. Le CV doit être compact pour tenir sur 1 seule page.\n\n"
            "RÈGLES DE LA LETTRE DE MOTIVATION :\n"
            "- La Lettre doit débuter par les coordonnées du candidat, la date, l'adresse de l'employeur, l'objet en gras (avec **), le corps du texte et la salutation (Cordialement, [Nom du Candidat]).\n\n"
            "FORMAT DE LA RÉPONSE :\n"
            "Génère d'abord le code HTML brut du CV. Pas de bloc ```html.\n"
            "Insère ensuite exactement le délimiteur :\n"
            "=== ANSCHREIBEN_START ===\n"
            "Génère ensuite la lettre de motivation en format texte brut."
        )

    user_prompt = (
        f"PROFIL DU CANDIDAT :\n{candidate_context}\n\n"
        f"OFFRE D'EMPLOI CIBLEE :\n{offer_context}\n\n"
        f"Date du jour : {today_str}\n\n"
        "Génère maintenant le CV canadien (HTML) et la lettre de motivation."
    )

    try:
        response = call_llm(system_prompt, user_prompt)
        if not response:
            return "", ""

        response = response.replace("```html", "").replace("```", "").strip()

        if "=== ANSCHREIBEN_START ===" in response:
            parts = response.split("=== ANSCHREIBEN_START ===")
            html_content = parts[0].strip()
            cover_letter = parts[1].strip()
        else:
            html_content = response
            cover_letter = ""

        return html_content, cover_letter
    except Exception as exc:
        log.error(f"Canada CV generation error: {exc}")
        return "", ""


@login_required
def improve_description_api(request):
    """
    Appelle Gemini pour réécrire la description de poste brute saisie par l'utilisateur
    au format canadien professionnel compatible ATS (verbes d'action, puces).
    """
    if request.method != "POST":
        return JsonResponse({"success": False, "error": "Méthode non autorisée."})

    import json
    try:
        data = json.loads(request.body)
        raw_text = data.get("text", "").strip()
    except Exception:
        raw_text = request.POST.get("text", "").strip()

    if not raw_text:
        return JsonResponse({"success": False, "error": "Veuillez saisir une description à améliorer."})

    from ai_engine.services.llm_service import call_llm

    system_prompt = (
        "Tu es un expert RH canadien. Transforme la description de tâche brute fournie "
        "en une liste à puces (bullet points) professionnelle de 3 à 5 éléments, en commençant par des verbes d'action à l'infinitif "
        "(ex: 'Gérer...', 'Assurer...', 'Collaborer...', 'Optimiser...'). Rends le contenu extrêmement professionnel "
        "et adapté pour passer les filtres ATS. Ne mets aucune introduction, salutation ou conclusion, commence directement par les puces."
    )

    try:
        improved_text = call_llm(system_prompt, f"Description brute :\n{raw_text}")
        if improved_text:
            return JsonResponse({"success": True, "improved_text": improved_text.strip()})
        return JsonResponse({"success": False, "error": "L'IA a retourné une réponse vide."})
    except Exception as e:
        log.error(f"Error in improve_description_api: {e}")
        return JsonResponse({"success": False, "error": str(e)})


@login_required
def immigration_diagnostic(request):
    """
    Diagnostic Express Entry : Calcule le score CRS et fournit une feuille de route
    personnalisée rédigée par l'IA (Gemini) en tant que consultant en immigration virtuelle.
    """
    profile, created = CanadaImmigrationProfile.objects.get_or_create(user=request.user)

    if request.method == "POST":
        form = CanadaImmigrationProfileForm(request.POST, instance=profile)
        if form.is_valid():
            p = form.save(commit=False)
            
            # Calcul du score CRS
            p.crs_score = _calculate_crs(
                p.age,
                p.education_level,
                p.work_experience_years,
                p.tcf_level,
                p.has_lmia_job
            )
            
            # Appel IA pour la feuille de route
            from ai_engine.services.llm_service import call_llm
            system_prompt = (
                "Tu es un consultant agréé en immigration pour le Canada (CRIC), bienveillant, rigoureux et stratégique.\n"
                "Analyse le profil du candidat, son score CRS obtenu sur l'Entrée Express et rédige une feuille de route "
                "sur mesure en français. Sois très constructif.\n"
                "RÈGLES DE CONSEIL :\n"
                "- Si le candidat n'a pas atteint le niveau C1/C2 (CLB 9+) en français (TCF), insiste sur le fait qu'il s'agit du levier le plus important : obtenir un C1/C2 lui donne un bonus massif de points (le bonus francophone Hors-Québec de 50 points + points de compétences transférables).\n"
                "- Encourage-le à s'entraîner sérieusement grâce aux leçons et examens blancs du TCF de notre plateforme.\n"
                "- Explique comment optimiser son profil (obtenir une EIMT via nos offres d'emploi Canada, passer l'anglais IELTS, etc.).\n"
                "Rends le texte clair et structuré en rubriques."
            )
            
            candidate_details = (
                f"Âge : {p.age} ans\n"
                f"Niveau d'études : {p.get_education_level_display()}\n"
                f"Expérience de travail à l'étranger : {p.work_experience_years} ans\n"
                f"Niveau de Français (TCF) : {p.get_tcf_level_display()}\n"
                f"Offre d'emploi approuvée par EIMT : {'Oui' if p.has_lmia_job else 'Non'}\n"
                f"Score CRS (Express Entry) calculé : {p.crs_score} points"
            )
            
            try:
                p.ai_roadmap = call_llm(system_prompt, candidate_details)
            except Exception as e:
                log.error(f"Error calling LLM for CRS roadmap: {e}")
                p.ai_roadmap = "Impossible de générer la feuille de route IA pour le moment. Veuillez réessayer."
                
            p.save()
            messages.success(request, "Votre diagnostic Express Entry IA a été mis à jour avec succès !")
            return redirect("canada_resume:diagnostic")
    else:
        form = CanadaImmigrationProfileForm(instance=profile)

    return render(request, "canada_resume/diagnostic.html", {
        "profile": profile,
        "form": form,
    })


def _calculate_crs(age, education, experience, french_level, has_lmia):
    score = 0
    
    # 1. Âge (maximum 110 points de 20 à 29 ans)
    if 20 <= age <= 29:
        score += 110
    elif age == 18:
        score += 99
    elif age == 19:
        score += 105
    elif age == 30:
        score += 105
    elif age == 31:
        score += 99
    elif age == 32:
        score += 94
    elif age == 33:
        score += 88
    elif age == 34:
        score += 83
    elif age == 35:
        score += 77
    elif age == 36:
        score += 72
    elif age == 37:
        score += 66
    elif age == 38:
        score += 61
    elif age == 39:
        score += 55
    elif age == 40:
        score += 50
    elif age == 41:
        score += 35
    elif age == 42:
        score += 25
    elif age == 43:
        score += 15
    elif age == 44:
        score += 5
    else:
        score += 0

    # 2. Scolarité (maximum 150 points)
    edu_points = {
        "doctorate": 150,
        "master": 135,
        "two_degrees": 128,
        "bachelor": 120,
        "two_year": 98,
        "one_year": 90,
        "high_school": 30,
    }
    score += edu_points.get(education, 30)

    # 3. Expérience professionnelle à l'étranger (maximum 25 points)
    if experience >= 3:
        score += 25
    elif experience == 2:
        score += 23
    elif experience == 1:
        score += 15
    else:
        score += 0

    # 4. Compétences linguistiques en Français (1ère langue officielle) + Bonus francophone
    # C1/C2 = CLB 9+. Donne de gros points de langue de base + bonus CRS Express Entry de 50 points.
    if french_level in ("C1", "C2"):
        score += 124  # Langue de base CLB 9/10
        score += 50   # Bonus francophone hors Québec
    elif french_level == "B2":
        score += 88   # Langue de base CLB 7/8
        score += 25   # Bonus partiel
    elif french_level == "B1":
        score += 48
    else:
        score += 0

    # 5. Offre d'emploi reservée EIMT
    if has_lmia:
        score += 50

    return score


def programs_hub(request):
    """Page d'accueil présentant tous les programmes d'immigration Canada (SEO Hub)."""
    return render(request, "canada_resume/programs_hub.html")


def program_ee(request):
    """Page explicative sur le programme Entrée Express (SEO)."""
    return render(request, "canada_resume/program_ee.html")


def program_arrima(request):
    """Page explicative sur le programme Arrima / Québec (SEO)."""
    return render(request, "canada_resume/program_arrima.html")


def program_pnp(request):
    """Page explicative sur le programme des Candidats des Provinces - PCP (SEO)."""
    return render(request, "canada_resume/program_pnp.html")


def program_others(request):
    """Page explicative sur les autres programmes d'immigration - AIP, RNIP, etc. (SEO)."""
    return render(request, "canada_resume/program_others.html")


def program_study(request):
    """Page explicative sur la procédure d'études au Canada, DLI et Bourses (SEO)."""
    return render(request, "canada_resume/program_study.html")


def program_visit(request):
    """Page explicative sur la procédure de visa tourisme/visiteur au Canada (SEO)."""
    return render(request, "canada_resume/program_visit.html")


from django.views.decorators.csrf import csrf_exempt

@login_required
def immigration_coach(request):
    """
    Page principale du coach IA pour l'immigration Canada.
    """
    is_pro = check_user_has_paid_edu_subscription(request.user)
    
    # Suivi des messages restants aujourd'hui pour les utilisateurs gratuits
    import datetime
    today_str = datetime.date.today().isoformat()
    session_date = request.session.get("canada_coach_limit_date")
    if session_date != today_str:
        request.session["canada_coach_limit_date"] = today_str
        request.session["canada_coach_message_count"] = 0
        
    messages_left = max(0, 5 - request.session.get("canada_coach_message_count", 0))
    
    return render(
        request,
        "canada_resume/immigration_coach.html",
        {
            "is_pro": is_pro,
            "messages_left": messages_left,
        }
    )


@login_required
@csrf_exempt
def immigration_coach_api(request):
    """
    API pour dialoguer avec le coach IA Immigration Canada.
    """
    is_pro = check_user_has_paid_edu_subscription(request.user)
    
    # Vérifier le quota quotidien pour les comptes gratuits
    if not is_pro:
        import datetime
        today_str = datetime.date.today().isoformat()
        session_date = request.session.get("canada_coach_limit_date")
        message_count = request.session.get("canada_coach_message_count", 0)
        
        if session_date != today_str:
            request.session["canada_coach_limit_date"] = today_str
            message_count = 0
            request.session["canada_coach_message_count"] = 0
            
        if message_count >= 5:
            return JsonResponse(
                {
                    "error": "subscription_required",
                    "reply": "Vous avez atteint votre limite gratuite de 5 messages par jour pour le Coach IA Immigration Canada. Veuillez vous abonner à E-Shelle Premium pour discuter en illimité !",
                },
                status=403,
            )
            
    if request.method != "POST":
        return JsonResponse({"error": "Only POST allowed"}, status=405)
        
    try:
        import json
        data = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)
        
    user_message = (data.get("message") or "").strip()
    history = data.get("history") or []
    
    if not user_message:
        return JsonResponse({"error": "Empty message"}, status=400)
        
    # Increment message count for free users
    if not is_pro:
        request.session["canada_coach_message_count"] = request.session.get("canada_coach_message_count", 0) + 1
        
    # Call LLM
    from ai_engine.services.llm_service import call_llm
    
    system_prompt = (
        "Tu es l'expert en immigration du Canada d'E-Shelle. Ton rôle est de conseiller "
        "les candidats à l'immigration sur toutes les procédures officielles basées sur le site "
        "du gouvernement du Canada (Canada.ca / IRCC). Réponds de façon précise, chaleureuse, professionnelle et structurée (utilises du gras et listes si besoin). "
        "Mentionne les programmes officiels : Entrée Express (FSTP, FSWP, CEC), Arrima (Québec), PNP (Candidats des Provinces), "
        "Permis d'études, PVT, ou Visa visiteur. Conseille toujours de faire évaluer son éligibilité avec "
        "notre Calculateur CRS E-Shelle."
    )
    
    # Format convo history
    prompt_builder = []
    for msg in history[-8:]: # last 8 messages
        role_label = "Candidat" if msg.get("role") == "user" else "Expert"
        prompt_builder.append(f"{role_label}: {msg.get('content')}")
    prompt_builder.append(f"Candidat: {user_message}\nExpert:")
    
    user_prompt = "\n".join(prompt_builder)
    
    try:
        reply = call_llm(system_prompt, user_prompt)
        return JsonResponse({"reply": reply.strip()})
    except Exception as e:
        return JsonResponse({"error": f"LLM error: {e}"}, status=500)


@login_required
def interview_simulation(request):
    """
    Page principale pour la simulation d'entretien canadien par IA.
    """
    is_pro = check_user_has_paid_edu_subscription(request.user)
    
    # Suivi des messages restants aujourd'hui pour les utilisateurs gratuits
    import datetime
    today_str = datetime.date.today().isoformat()
    session_date = request.session.get("canada_interview_limit_date")
    if session_date != today_str:
        request.session["canada_interview_limit_date"] = today_str
        request.session["canada_interview_message_count"] = 0
        
    messages_left = max(0, 5 - request.session.get("canada_interview_message_count", 0))
    
    return render(
        request,
        "canada_resume/interview_simulation.html",
        {
            "is_pro": is_pro,
            "messages_left": messages_left,
        }
    )


@login_required
@csrf_exempt
def interview_simulation_api(request):
    """
    API pour dialoguer avec l'interviewer IA (simulation d'entretien).
    """
    is_pro = check_user_has_paid_edu_subscription(request.user)
    
    # Vérifier le quota quotidien pour les comptes gratuits
    if not is_pro:
        import datetime
        today_str = datetime.date.today().isoformat()
        session_date = request.session.get("canada_interview_limit_date")
        message_count = request.session.get("canada_interview_message_count", 0)
        
        if session_date != today_str:
            request.session["canada_interview_limit_date"] = today_str
            message_count = 0
            request.session["canada_interview_message_count"] = 0
            
        if message_count >= 5:
            return JsonResponse(
                {
                    "error": "subscription_required",
                    "reply": "Vous avez atteint votre limite gratuite de 5 messages par jour pour la simulation d'entretien. Veuillez vous abonner à E-Shelle Premium pour vous entraîner en illimité !",
                },
                status=403,
            )
            
    if request.method != "POST":
        return JsonResponse({"error": "Only POST allowed"}, status=405)
        
    try:
        import json
        data = json.loads(request.body.decode("utf-8"))
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)
        
    user_message = (data.get("message") or "").strip()
    history = data.get("history") or []
    sector = (data.get("sector") or "Général").strip()
    
    if not user_message and len(history) > 0:
        return JsonResponse({"error": "Empty message"}, status=400)
        
    # Increment message count for free users
    if not is_pro and user_message:
        request.session["canada_interview_message_count"] = request.session.get("canada_interview_message_count", 0) + 1
        
    from ai_engine.services.llm_service import call_llm
    
    system_prompt = (
        f"Tu es un recruteur de ressources humaines (RH) pour une entreprise canadienne qui mène un entretien d'embauche pour le secteur : '{sector}'.\n"
        "RÈGLES CRITIQUES DE L'ENTRETIEN CANADIEN :\n"
        "- Ton but est d'évaluer le candidat de façon constructive et de lui poser des questions de recrutement standard (comportementales ou techniques).\n"
        "- Pour CHAQUE réponse envoyée par le candidat, fournis d'abord une courte ÉVALUATION constructive de sa réponse (ex: 'Bonne réponse, mais vous devriez quantifier vos résultats', ou 'Excellente utilisation de la méthode STAR') puis pose la QUESTION SUIVANTE.\n"
        "- Pose UNE SEULE question à la fois.\n"
        "- Reste professionnel, courtois, encourageant et formule tes retours selon les standards RH canadiens."
    )
    
    if not user_message and len(history) == 0:
        # First message: initiate interview
        try:
            reply = call_llm(
                system_prompt, 
                f"Débute l'entretien pour le secteur {sector}. Salue chaleureusement le candidat et pose la première question."
            )
            return JsonResponse({"reply": reply.strip()})
        except Exception as e:
            return JsonResponse({"error": f"LLM error: {e}"}, status=500)
            
    # Format convo history
    prompt_builder = []
    for msg in history[-8:]:
        role_label = "Candidat" if msg.get("role") == "user" else "Recruteur"
        prompt_builder.append(f"{role_label}: {msg.get('content')}")
    prompt_builder.append(f"Candidat: {user_message}\nRecruteur:")
    
    user_prompt = "\n".join(prompt_builder)
    
    try:
        reply = call_llm(system_prompt, user_prompt)
        return JsonResponse({"reply": reply.strip()})
    except Exception as e:
        return JsonResponse({"error": f"LLM error: {e}"}, status=500)


def talents_directory(request):
    """
    Annuaire public des profils de talents / candidats Canada pour les employeurs.
    """
    # Récupérer tous les profils de candidats complets
    talents = CanadaCVProfile.objects.exclude(first_name="").order_by("-updated_at")
    
    # Filtrer par secteur si spécifié
    sector = request.GET.get("sector", "").strip()
    if sector:
        talents = talents.filter(target_sector__icontains=sector)
        
    # Liste des secteurs pour le filtre
    sectors = CanadaCVProfile.objects.exclude(target_sector="").values_list("target_sector", flat=True).distinct()
    
    return render(
        request,
        "canada_resume/talents.html",
        {
            "talents": talents,
            "sectors": sectors,
            "selected_sector": sector,
            "is_pro": check_user_has_paid_edu_subscription(request.user) if request.user.is_authenticated else False,
        }
    )


@login_required
def manage_projects(request):
    if request.method == "POST":
        form = CanadaCVProjectForm(request.POST, request.FILES)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.user = request.user
            obj.save()
            messages.success(request, "Projet / Réalisation ajouté(e) au portfolio !")
            return redirect("canada_resume:dashboard")
    else:
        form = CanadaCVProjectForm()

    projects = CanadaCVProject.objects.filter(user=request.user)
    return render(request, "canada_resume/manage_projects.html", {
        "form": form, "projects": projects
    })


@login_required
def delete_project(request, pk):
    project = get_object_or_404(CanadaCVProject, pk=pk, user=request.user)
    project.delete()
    messages.success(request, "Réalisation supprimée du portfolio.")
    return redirect("canada_resume:dashboard")


def talent_detail(request, pk):
    """
    Fiche détaillée publique d'un candidat pour les recruteurs.
    """
    talent = get_object_or_404(CanadaCVProfile, pk=pk)
    # Fetch all details for this candidate
    experiences = CanadaCVExperience.objects.filter(user=talent.user)
    educations = CanadaCVEducation.objects.filter(user=talent.user)
    languages = CanadaCVLanguage.objects.filter(user=talent.user)
    projects = CanadaCVProject.objects.filter(user=talent.user)
    
    is_pro = check_user_has_paid_edu_subscription(request.user) if request.user.is_authenticated else False
    
    return render(
        request,
        "canada_resume/talent_detail.html",
        {
            "talent": talent,
            "experiences": experiences,
            "educations": educations,
            "languages": languages,
            "projects": projects,
            "is_pro": is_pro,
        }
    )


def canada_resources(request):
    """
    Liste des ressources d'immigration Canada (PDF, Vidéos, Excel).
    """
    is_pro = check_user_has_paid_edu_subscription(request.user) if request.user.is_authenticated else False
    
    resources = CanadaResource.objects.filter(is_active=True)
    
    # Filtre par type
    rtype = request.GET.get("type", "").strip()
    if rtype:
        resources = resources.filter(resource_type=rtype)
        
    # Recherche
    q = request.GET.get("q", "").strip()
    if q:
        from django.db.models import Q
        resources = resources.filter(
            Q(title__icontains=q) |
            Q(description__icontains=q)
        )
        
    return render(
        request,
        "canada_resume/resources.html",
        {
            "resources": resources,
            "selected_type": rtype,
            "q": q,
            "is_pro": is_pro,
        }
    )


@login_required
def download_resource(request, pk):
    """
    Télécharge ou visionne une ressource avec vérification Premium.
    """
    res = get_object_or_404(CanadaResource, pk=pk, is_active=True)
    is_pro = check_user_has_paid_edu_subscription(request.user)
    
    if res.is_premium and not is_pro:
        messages.warning(request, "Cette ressource est réservée aux membres Canada Premium. Veuillez souscrire à un abonnement.")
        return redirect(reverse("accounts:upgrade") + "?app=prep")
        
    # Incrémenter le compteur
    CanadaResource.objects.filter(pk=res.pk).update(downloads_count=res.downloads_count + 1)
    
    if res.resource_type == "video" and res.video_url:
        return redirect(res.video_url)
    elif res.file:
        return redirect(res.file.url)
        
    messages.error(request, "Fichier ou lien vidéo non trouvé pour cette ressource.")
    return redirect("canada_resume:canada_resources")


