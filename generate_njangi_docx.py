import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def add_heading_with_bottom_border(doc, text, level):
    h = doc.add_heading(text, level=level)
    hPr = h._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # 1.5 pt
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '0D3B6E')
    pBdr.append(bottom)
    hPr.append(pBdr)
    return h

def generate_docx(output_path):
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run("NJANGI+")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(38)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0D, 0x3B, 0x6E)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(40)
    sub_run = subtitle_p.add_run("La tontine numérique révolutionnaire de la suite E-Shelle\nDossier de Présentation Complet & Spécifications")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_before = Pt(150)
    info_run = info_p.add_run(
        "Préparé par l'équipe E-Shelle Suite App\n"
        "À l'attention des parents, membres et investisseurs potentiels\n"
        "Version 1.0 — Juillet 2026\n"
        "Contact : contact@e-shelle.com"
    )
    info_run.font.name = 'Arial'
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    
    doc.add_page_break()
    
    # ---------------------------------------------------------
    # SECTION 1: INTRODUCTION
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "1. Contexte : La Tontine traditionnelle et ses limites", level=1)
    doc.add_paragraph(
        "En Afrique subsaharienne et au sein de ses diasporas, la tontine (ou Njangui) est bien plus qu'une association d'épargne. "
        "Elle représente le cœur financier et social de la communauté. Elle permet d'accumuler du capital sans passer par les banques traditionnelles, "
        "d'octroyer des crédits à taux préférentiels, et de soutenir les membres lors d'événements familiaux (mariages, naissances, deuils)."
    )
    doc.add_paragraph(
        "Cependant, la gestion traditionnelle des tontines fait face à plusieurs limites opérationnelles majeures :"
    )
    
    doc.add_paragraph("• Perte de confiance et opacité : Les cahiers physiques et fichiers Excel sont vulnérables à la perte, au vol ou aux falsifications volontaires.", style='List Bullet')
    doc.add_paragraph("• Risques liés au cash physique : La manipulation de fortes sommes d'argent liquide pendant les réunions présente des risques réels de vol, d'agression ou d'erreurs de caisse.", style='List Bullet')
    doc.add_paragraph("• Complexité administrative : Le calcul pro-rata des intérêts de l'épargne, les reliquats de prêts restants, et les pénalités de retard prennent des heures de calcul manuel à chaque fin de mois, décourageant les gestionnaires.", style='List Bullet')
    doc.add_paragraph("• Exclusion de la diaspora : Les membres vivant à l'étranger peinent à suivre l'évolution des cotisations et à envoyer ou recevoir des fonds en temps voulu sans intermédiaires douteux.", style='List Bullet')
    
    # ---------------------------------------------------------
    # SECTION 2: LA VISION NJANGI+
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "2. La vision de Njangi+ : Confiance et Rigueur", level=1)
    doc.add_paragraph(
        "Njangi+ transpose la dynamique de confiance des tontines traditionnelles dans un environnement numérique sécurisé et moderne. "
        "L'application a été conçue pour s'adapter à la réalité locale du terrain, et non l'inverse. C'est pourquoi elle repose sur deux principes essentiels :"
    )
    
    doc.add_paragraph("1. Saisie exclusive par le Bureau : Pour éviter que les membres fassent de fausses déclarations de paiement en ligne ou perturbent la comptabilité générale, l'ensemble des opérations financières (dépôts, remboursements, cotisations, fonds de caisse) est obligatoirement saisi par le Bureau (Président, Trésorier ou Secrétaire) lors des séances physiques.", style='List Bullet')
    doc.add_paragraph("2. Transparence totale pour les membres : Chaque membre simple dispose d'un espace personnel sécurisé en ligne où il suit instantanément l'évolution de ses avoirs, de ses prêts et de ses relevés, en toute transparence et sans action technique requise de sa part.", style='List Bullet')
    
    # ---------------------------------------------------------
    # SECTION 3: ESPACE BUREAU
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "3. Fonctionnalités de l'Espace Bureau", level=1)
    doc.add_paragraph(
        "L'Espace Bureau est le centre névralgique de Njangi+. Il permet aux administrateurs de piloter les réunions physiques de manière fluide et professionnelle :"
    )
    
    doc.add_paragraph("• Gestion des séances de réunion : Planification, enregistrement automatique de la fiche de présence et génération instantanée des cotisations ordinaires attendues pour la séance.", style='List Bullet')
    
    doc.add_paragraph("• Bénéficiaires multiples (Mains levées) : Contrairement aux systèmes rigides qui n'autorisent qu'un seul gagnant, Njangi+ permet d'attribuer la main levée de la séance à un ou plusieurs bénéficiaires en simultané. La plateforme calcule automatiquement la répartition financière correspondante.", style='List Bullet')
    
    doc.add_paragraph("• Gestion avancée du Fond de Caisse (Fonds de base) : Chaque membre dispose d'un fond de caisse obligatoire. Lors des séances, le bureau peut y enregistrer des versements. Le système affiche automatiquement à l'écran le solde en direct du membre sélectionné pour guider le trésorier.", style='List Bullet')
    
    doc.add_paragraph("• Retraits collectifs et individuels de fond de caisse : En cas d'événement malheureux (deuil) ou heureux (mariage) affectant la tontine, le bureau peut effectuer un prélèvement individuel ou collectif sur le fond de caisse de tous les membres simultanément, avec saisie obligatoire d'un motif.", style='List Bullet')
    
    doc.add_paragraph("• Gestion des membres virtuels (Inclusion des mamans) : Pour inclure les personnes âgées ou n'ayant pas de compte e-mail/smartphone, le bureau peut créer des comptes de membres virtuels. Le système gère ces comptes exactement comme les autres dans les rapports et les flux comptables, le bureau agissant en tant que tuteur.", style='List Bullet')
    
    doc.add_paragraph("• Bouton d'annulation (Rollback) : Toutes les actions comptables enregistrées en séance possèdent un bouton d'annulation rapide pour corriger instantanément une erreur de saisie du bureau.", style='List Bullet')
    
    # ---------------------------------------------------------
    # SECTION 4: ESPACE MEMBRE
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "4. Fonctionnalités de l'Espace Membre (Consultatif)", level=1)
    doc.add_paragraph(
        "L'espace membre a été simplifié pour être intuitif et consultatif, éliminant tout stress technique pour les utilisateurs :"
    )
    
    doc.add_paragraph("• Tableau de bord d'ensemble : Un récapitulatif graphique de l'épargne cumulée, des prêts en cours et du score de fiabilité de remboursement.", style='List Bullet')
    doc.add_paragraph("• Suivi des prêts et échéances : Visualisation du taux d'intérêt appliqué, de la date limite de remboursement et du reste dû exact.", style='List Bullet')
    doc.add_paragraph("• Suivi des cotisations et du fond de caisse : Historique clair et daté de toutes les cotisations versées et des retraits appliqués au fond de caisse.", style='List Bullet')
    doc.add_paragraph("• Relevé PDF de situation : Génération instantanée d'un relevé d'intérêts certifié listant les avoirs du membre et sa part de distribution dans le pool d'intérêts.", style='List Bullet')
    
    # ---------------------------------------------------------
    # SECTION 5: COMPTABILITÉ & SÉCURITÉ
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "5. Comptabilité et Sécurité financière", level=1)
    doc.add_paragraph(
        "Njangi+ implémente des algorithmes de niveau bancaire adaptés aux usages de la tontine :"
    )
    
    doc.add_paragraph("• Réserve de sécurité obligatoire : Par défaut, la plateforme applique une réserve de 20% sur la caisse totale du fond commun. Si le fond commun contient 100 000 FCFA, la tontine n'autorise des prêts qu'à hauteur de 80 000 FCFA pour garantir un fond de roulement minimal de sécurité.", style='List Bullet')
    
    doc.add_paragraph("• Journal d'audit (Audit Trail) : Chaque action administrative (création de prêt, retrait de fond de caisse, modification de rôle) est enregistrée dans un journal d'audit immuable. Le président et les membres du bureau peuvent ainsi retracer l'origine de chaque modification (qui, quand, quoi) pour exclure toute fraude.", style='List Bullet')
    
    doc.add_paragraph("• Distribution d'intérêts au prorata : À la fin de chaque mois, les intérêts cumulés sur les remboursements de prêts sont distribués automatiquement aux membres épargnants au prorata de leurs dépôts actifs, récompensant équitablement l'épargne de chacun.", style='List Bullet')
    
    # ---------------------------------------------------------
    # SECTION 6: MODÈLE PREMIUM
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "6. Modèle économique", level=1)
    doc.add_paragraph(
        "Pour assurer le support et la maintenance continue, Njangi+ propose des plans d'abonnement clairs et abordables :"
    )
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Plan d\'abonnement'
    hdr_cells[1].text = 'Tarif mensuel'
    hdr_cells[2].text = 'Fonctionnalités incluses'
    
    for cell in hdr_cells:
        set_cell_background(cell, "0D3B6E")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.bold = True
        
    plans = [
        ("Plan Gratuit (Free)", "0 FCFA", "Jusqu'à 5 membres, gestion de base des cotisations."),
        ("Plan Standard", "3 000 FCFA / groupe", "Membres illimités, gestion des séances et des prêts."),
        ("Plan Pro", "7 000 FCFA / groupe", "Ajoute la piste d'audit, les membres virtuels et les relevés PDF."),
        ("Plan Association", "15 000 FCFA / groupe", "Multi-bureau, gestion de plusieurs tontines simultanées, support VIP.")
    ]
    
    for i, (name, price, desc) in enumerate(plans):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = name
        row_cells[1].text = price
        row_cells[2].text = desc
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        
    doc.add_paragraph("") # Spacer
    
    # ---------------------------------------------------------
    # SECTION 7: PERSPECTIVES ET PERSPECTIVES D'AVENIR
    # ---------------------------------------------------------
    add_heading_with_bottom_border(doc, "7. Perspectives d'avenir : Intégration Fintech", level=1)
    doc.add_paragraph(
        "L'évolution naturelle de Njangi+ se fera par des partenariats stratégiques de paiement :"
    )
    doc.add_paragraph("• Intégration des APIs de Mobile Money (MTN MoMo, Orange Money) pour automatiser la réception et le versement des fonds.", style='List Bullet')
    doc.add_paragraph("• Attribution d'historique de crédit (Credit Scoring) basé sur la régularité et le score de fiabilité de l'application pour aider les membres à obtenir des microcrédits bancaires classiques.", style='List Bullet')
    doc.add_paragraph("• Extension vers le e-commerce communautaire pour proposer des achats groupés de matières premières agricoles ou de gaz domestique (E-Shelle Gaz) directement financés par la tontine.", style='List Bullet')
    
    doc.add_paragraph("\nEn conclusion, Njangi+ associe la solidarité de la tradition africaine à la rigueur des outils numériques modernes pour créer la tontine de demain. C'est un outil indispensable pour valoriser et sécuriser l'épargne communautaire.")
    
    doc.save(output_path)

if __name__ == "__main__":
    import sys
    output_file = "Njangi_Presentation_Complete.docx"
    if len(sys.argv) > 1:
        output_file = sys.argv[1]
    generate_docx(output_file)
    print(f"Word document generated at {output_file}")
