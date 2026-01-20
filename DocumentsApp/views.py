from io import BytesIO
import os
import tempfile
import subprocess
from django.http import FileResponse
from .forms import ConversionForm

from django.contrib import messages
from django.core.files.base import ContentFile
from django.shortcuts import render, redirect, get_object_or_404
from django.utils.text import slugify

from docx import Document
from deep_translator import GoogleTranslator
from pypdf import PdfReader, PdfWriter

from .forms import TranslationForm, CompressionForm
from .models import TranslationJob


def documents_home(request):
    return render(request, "DocumentsApp/home.html")


import time
import traceback

def safe_translate(translator, text: str, max_retries: int = 2) -> str:
    """
    Traduction robuste :
    - ignore textes vides
    - retry en cas de soucis
    - retourne le texte original si échec
    """
    if not text:
        return text

    # évite de traduire des “tokens” inutiles (emails, numéros simples)
    stripped = text.strip()
    if not stripped:
        return text

    for _ in range(max_retries + 1):
        try:
            out = translator.translate(stripped)

            # deep_translator peut renvoyer None parfois
            if out is None:
                return text

            # sécurité : on force string
            return str(out)

        except Exception:
            time.sleep(0.3)

    return text


import time
import re
from io import BytesIO
from docx import Document
from deep_translator import GoogleTranslator

# Marqueurs très rares (pour survivre à la traduction)
MARK_OPEN = "⟪IMM97_RUN_"
MARK_CLOSE = "⟫"
SEP = " ⟪IMM97_SEP⟫ "  # séparateur de paragraphes (si tu batch plus tard)

def _safe_translate(translator: GoogleTranslator, text: str, retries: int = 2) -> str:
    """
    Traduction robuste:
    - retry en cas d'erreur réseau/rate limit
    - si None => on renvoie le texte original
    """
    last_err = None
    for _ in range(retries + 1):
        try:
            out = translator.translate(text)
            if out is None:
                return text
            return out
        except Exception as e:
            last_err = e
            time.sleep(0.6)
    # si tout échoue, on renvoie le texte original (et on laissera la view afficher l'erreur)
    raise last_err

def _translate_paragraph_preserve_runs(paragraph, translator: GoogleTranslator):
    """
    Améliore la qualité (contexte) tout en gardant la mise en forme extrême (runs).
    Principe:
    - on concatène les runs avec des marqueurs invariants
    - on traduit la chaîne complète (le traducteur voit le contexte global)
    - on re-split par marqueurs et on réécrit run.text (format inchangé)
    """
    if not paragraph.runs:
        return

    # Construire une chaîne avec marqueurs run par run
    parts = []
    for i, run in enumerate(paragraph.runs):
        t = run.text
        if t is None:
            t = ""
        # On garde même les runs vides (pour ne pas casser la structure)
        parts.append(f"{MARK_OPEN}{i}{MARK_CLOSE}{t}{MARK_OPEN}/{i}{MARK_CLOSE}")

    payload = "".join(parts)

    # Si tout est vide/espaces, on skip
    if not payload.strip():
        return

    translated = _safe_translate(translator, payload)

    # Extraire chaque segment traduit entre marqueurs
    # Exemple: ⟪IMM97_RUN_0⟫ ... ⟪IMM97_RUN_/0⟫
    segments = {}
    for i in range(len(paragraph.runs)):
        pattern = re.escape(f"{MARK_OPEN}{i}{MARK_CLOSE}") + r"(.*?)" + re.escape(f"{MARK_OPEN}/{i}{MARK_CLOSE}")
        m = re.search(pattern, translated, flags=re.DOTALL)
        if not m:
            segments = None
            break
        segments[i] = m.group(1)

    # Fallback: si les marqueurs ont été abîmés, on traduit tout le paragraphe et on met dans le 1er run
    if segments is None:
        full_text = paragraph.text or ""
        if full_text.strip():
            best = _safe_translate(translator, full_text)
            # on garde la forme du 1er run, on vide les autres
            paragraph.runs[0].text = best
            for r in paragraph.runs[1:]:
                r.text = ""
        return

    # Réinjecter dans les runs (format inchangé)
    for i, run in enumerate(paragraph.runs):
        run.text = segments.get(i, "")

def translate_docx_file(uploaded_file, source_lang, target_lang):
    """
    Traduction DOCX haute qualité + conservation extrême mise en forme.
    - Contextualise la traduction (par paragraphe)
    - Préserve la mise en forme (runs + tableaux)
    """
    uploaded_file.seek(0)
    buffer_in = BytesIO(uploaded_file.read())
    doc = Document(buffer_in)

    translator = GoogleTranslator(
        source="auto" if source_lang == "auto" else source_lang,
        target=target_lang
    )

    # Paragraphes normaux
    for paragraph in doc.paragraphs:
        # éviter les paragraphes vides
        if (paragraph.text or "").strip():
            _translate_paragraph_preserve_runs(paragraph, translator)

    # Tableaux (très fréquent sur diplômes/relevés)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if (paragraph.text or "").strip():
                        _translate_paragraph_preserve_runs(paragraph, translator)

    buffer_out = BytesIO()
    doc.save(buffer_out)
    buffer_out.seek(0)
    return buffer_out



def translation(request):
    form = TranslationForm(request.POST or None, request.FILES or None)

    if request.method == "POST" and form.is_valid():
        uploaded_file = form.cleaned_data["file"]
        source_lang = form.cleaned_data["source_lang"]
        target_lang = form.cleaned_data["target_lang"]

        job = TranslationJob.objects.create(
            original_name=uploaded_file.name,
            source_lang=source_lang,
            target_lang=target_lang,
            original_file=uploaded_file,
        )

        try:
            translated_io = translate_docx_file(job.original_file, source_lang, target_lang)

            base = uploaded_file.name.rsplit(".", 1)[0]
            safe = slugify(base)[:80] or "document"
            output_name = f"{safe}_traduit_{target_lang}.docx"

            job.translated_file.save(output_name, ContentFile(translated_io.getvalue()), save=True)

            messages.success(request, "Traduction terminée. Ton fichier est prêt.")
            return redirect("DocumentsApp:translation_result", pk=job.pk)

        except Exception as e:
            messages.error(request, f"Erreur lors de la traduction : {e}")
            return redirect("DocumentsApp:translation")

    return render(request, "DocumentsApp/translation_form.html", {"form": form})


def translation_result(request, pk):
    result = get_object_or_404(TranslationJob, pk=pk)
    if not result.translated_file:
        messages.error(request, "Le fichier traduit n'est pas disponible.")
        return redirect("DocumentsApp:translation")
    return render(request, "DocumentsApp/translation_result.html", {"result": result})


def compress_pdf_file(uploaded_file):
    uploaded_file.seek(0)
    buffer_in = BytesIO(uploaded_file.read())
    reader = PdfReader(buffer_in)
    writer = PdfWriter()

    for page in reader.pages:
        writer.add_page(page)

    writer.add_metadata(reader.metadata or {})

    buffer_out = BytesIO()
    writer.write(buffer_out)
    buffer_out.seek(0)
    return buffer_out


def compression(request):
    from django.http import FileResponse

    form = CompressionForm(request.POST or None, request.FILES or None)

    if request.method == "POST" and form.is_valid():
        uploaded_file = form.cleaned_data["file"]

        try:
            compressed_io = compress_pdf_file(uploaded_file)
            base_name = uploaded_file.name.rsplit(".", 1)[0]
            output_name = f"{base_name}_compressed.pdf"

            return FileResponse(compressed_io, as_attachment=True, filename=output_name)

        except Exception as e:
            messages.error(request, f"Erreur lors de la compression : {e}")
            return redirect("DocumentsApp:compression")

    return render(request, "DocumentsApp/compression.html", {"form": form})

def _find_libreoffice():
    """
    Retourne un chemin vers soffice si dispo (Windows/Linux).
    """
    candidates = [
        os.environ.get("LIBREOFFICE_PATH"),
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
        "libreoffice",
    ]
    for c in candidates:
        if not c:
            continue
        # si c'est un chemin windows
        if c.endswith(".exe") and os.path.exists(c):
            return c
        # sinon commande système
        return c
    return None


def convert_docx_to_pdf(uploaded_file) -> BytesIO:
    """
    Conversion DOCX -> PDF
    Priorité: LibreOffice headless (meilleur en prod).
    Fallback Windows: docx2pdf si Word installé.
    """
    uploaded_file.seek(0)

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "input.docx")
        out_path = os.path.join(tmpdir, "input.pdf")

        with open(in_path, "wb") as f:
            f.write(uploaded_file.read())

        soffice = _find_libreoffice()

        # 1) LibreOffice (recommandé)
        if soffice:
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, in_path],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                if not os.path.exists(out_path):
                    # LibreOffice peut nommer différemment; on cherche le premier PDF
                    pdfs = [p for p in os.listdir(tmpdir) if p.lower().endswith(".pdf")]
                    if pdfs:
                        out_path = os.path.join(tmpdir, pdfs[0])

                with open(out_path, "rb") as f:
                    return BytesIO(f.read())

            except Exception:
                pass  # fallback plus bas

        # 2) Fallback Windows: docx2pdf (Word requis)
        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(in_path, out_path)
            with open(out_path, "rb") as f:
                return BytesIO(f.read())
        except Exception as e:
            raise RuntimeError(
                "Conversion DOCX→PDF impossible. "
                "Installe LibreOffice (recommandé) ou Microsoft Word (docx2pdf). "
                f"Détails: {e}"
            )


def convert_pdf_to_docx(uploaded_file) -> BytesIO:
    """
    Conversion PDF -> DOCX
    Bonne qualité si PDF texte. Si PDF scanné: résultat faible (OCR requis).
    """
    uploaded_file.seek(0)

    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = os.path.join(tmpdir, "input.pdf")
        out_path = os.path.join(tmpdir, "output.docx")

        with open(in_path, "wb") as f:
            f.write(uploaded_file.read())

        try:
            from pdf2docx import Converter
            cv = Converter(in_path)
            cv.convert(out_path, start=0, end=None)
            cv.close()

            with open(out_path, "rb") as f:
                return BytesIO(f.read())

        except Exception as e:
            raise RuntimeError(
                "Conversion PDF→DOCX impossible. "
                "Assure-toi que le PDF contient du texte (pas un scan). "
                f"Détails: {e}"
            )


def conversion(request):
    """
    Page conversion (DOCX⇄PDF) + téléchargement du fichier converti.
    """
    form = ConversionForm(request.POST or None, request.FILES or None)

    if request.method == "POST" and form.is_valid():
        mode = form.cleaned_data["mode"]
        uploaded_file = form.cleaned_data["file"]

        try:
            if mode == "docx_to_pdf":
                out_io = convert_docx_to_pdf(uploaded_file)
                base = uploaded_file.name.rsplit(".", 1)[0]
                filename = f"{base}.pdf"
                out_io.seek(0)
                return FileResponse(out_io, as_attachment=True, filename=filename)

            if mode == "pdf_to_docx":
                out_io = convert_pdf_to_docx(uploaded_file)
                base = uploaded_file.name.rsplit(".", 1)[0]
                filename = f"{base}.docx"
                out_io.seek(0)
                return FileResponse(out_io, as_attachment=True, filename=filename)

            messages.error(request, "Type de conversion invalide.")
            return redirect("DocumentsApp:conversion")

        except Exception as e:
            messages.error(request, f"Erreur lors de la conversion : {e}")
            return redirect("DocumentsApp:conversion")

    return render(request, "DocumentsApp/conversion.html", {"form": form})
